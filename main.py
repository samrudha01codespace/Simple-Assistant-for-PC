import pickle  # Used for serializing and deserializing Python objects
import threading  # Used for running tasks concurrently in separate threads
import time  # Provides time-related functions such as sleep and getting the current time
import tkinter as tk  # Provides tools for creating graphical user interfaces
import webbrowser  # Allows opening URLs in a web browser
from datetime import datetime  # Provides functions for manipulating dates and times
from io import BytesIO  # Provides a buffer for in-memory byte streams
from tkinter import filedialog, simpledialog  # Provides tools for opening file dialogs in a tkinter GUI
import openrouteservice  # Provides access to the OpenRouteService API for mapping and routing services
import os  # Provides functions for interacting with the operating system
import pygame  # A library for creating games and multimedia applications
import cv2  # OpenCV library for computer vision tasks like image processing and video capture
import folium  # Used for creating interactive maps
import google.generativeai as genai  # Provides access to Google Generative AI services
import matplotlib.pyplot as plt  # A plotting library for creating static, animated, and interactive visualizations
import pyjokes  # A library for generating programming-related jokes
import pyttsx3  # A text-to-speech conversion library
import requests  # Allows sending HTTP requests to interact with web services
import speech_recognition as sr  # A library for recognizing speech and converting it to text
from PIL import Image  # Python Imaging Library for opening, manipulating, and saving image files
from docx import Document  # Used for creating, reading, and editing Microsoft Word (.docx) files

os.environ['TF_ENABLE_ONEDNN_OPTS'] = '0'

from transformers import GPT2LMHeadModel, GPT2Tokenizer

# Load pre-trained model and tokenizer
model_name = "gpt2"
tokenizer = GPT2Tokenizer.from_pretrained(model_name)
model = GPT2LMHeadModel.from_pretrained(model_name)


def generate_response(prompt):
    inputs = tokenizer.encode(prompt, return_tensors="pt")
    outputs = model.generate(inputs, max_length=150, num_return_sequences=1, no_repeat_ngram_size=2)
    response = tokenizer.decode(outputs[0], skip_special_tokens=True)
    return response


# Initialize pyttsx3 engine for TTS
engine = pyttsx3.init()

NEWS_API_KEY = ""
ORS_API_KEY = ""  # Replace with your actual OpenRouteService API key
GOOGLE_API_KEY = ""  # Replace with your actual Google API key
CSE_ID = ""
WEATHER_API_KEY = ""  # Add your OpenWeatherMap API key
NASA_API_KEY = ""

data_folder = "data"  # Path to your "Data" folder


# to play music add the folder named "data" in that add another folder named "music" in that add musics
def play_music(partial_name):
    pygame.init()
    pygame.mixer.init()

    music_folder = "data/music"
    found_music = []

    partial_name_lower = partial_name.lower()

    for filename in os.listdir(music_folder):
        if partial_name_lower in filename.lower():
            found_music.append(filename)

    if found_music:
        if len(found_music) == 1:
            music_path = os.path.join(music_folder, found_music[0])
            play_audio("Do you want to play the music in loop or once?")
            loop_choice = get_voice_input().lower()
            if "continuous" in loop_choice:
                pygame.mixer.music.load(music_path)
                pygame.mixer.music.play(-1)  # -1 for infinite loop
                print(f"Playing music in loop")
                stop_music_thread = threading.Thread(target=stop_music_listener)
                stop_music_thread.start()
            else:
                pygame.mixer.music.load(music_path)
                pygame.mixer.music.play()
                print(f"Playing {found_music[0]} once")
        else:
            print("Found matching music files:")
            for idx, music_file in enumerate(found_music, 1):
                print(f"{idx}. {music_file}")

            selection = get_voice_input()

            try:
                selection_idx = int(selection) - 1
                if 0 <= selection_idx < len(found_music):
                    music_path = os.path.join(music_folder, found_music[selection_idx])
                    play_audio("Do you want to play the music continuous or once?")
                    loop_choice = get_voice_input().lower()
                    if "loop" in loop_choice:
                        pygame.mixer.music.load(music_path)
                        pygame.mixer.music.play(-1)  # -1 for infinite loop
                        print(f"Playing {found_music[selection_idx]} in loop")
                        stop_music_thread = threading.Thread(target=stop_music_listener)
                        stop_music_thread.daemon = True
                        stop_music_thread.start()

                    else:
                        pygame.mixer.music.load(music_path)
                        pygame.mixer.music.play()
                        print(f"Playing {found_music[selection_idx]} once")
                else:
                    print("Invalid selection. Please try again.")
            except ValueError:
                print("Invalid input. Please enter a number.")
    else:
        print("No matching music files found.")


# Example usage:


# Function to play audio using pyttsx3
def play_audio(text):
    engine.say(text)
    engine.runAndWait()


# Function to locate realtime location


# Function to get voice input
def get_voice_input():
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        recognizer.adjust_for_ambient_noise(source)
        print("Please speak your query...")
        audio = recognizer.listen(source)
        try:
            query = recognizer.recognize_google(audio)
            print("You said:", query)
            return query
        except sr.UnknownValueError:
            print("Sorry, I could not understand your query.")
        except sr.RequestError:
            print("Could not request results; check your network connection.")


def wake_word_listener():
    while True:
        print("Listening for wake word...")
        query = get_voice_input()
        if query and "jarvis" in query:
            play_audio("Yes Sir")
            return


# Function to recognize fa  ces using OpenCV
def recognize_faces():
    face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
    video_capture = cv2.VideoCapture(0)

    face_dir = "faces"
    if not os.path.exists(face_dir):
        os.mkdir(face_dir)

    known_faces = {}
    if os.path.exists("known_faces.pkl"):
        with open("known_faces.pkl", "rb") as f:
            known_faces = pickle.load(f)

    recognized_name = None

    while True:
        print(cv2.__version__)
        ret, frame = video_capture.read()
        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
        faces = face_cascade.detectMultiScale(gray, scaleFactor=1.1, minNeighbors=5, minSize=(30, 30))

        for (x, y, w, h) in faces:
            face_image = gray[y:y + h, x:x + w]
            recognized = False
            for name, saved_face in known_faces.items():
                result = cv2.matchTemplate(face_image, saved_face, cv2.TM_CCOEFF_NORMED)
                _, max_val, _, _ = cv2.minMaxLoc(result)
                if max_val > 0.6:
                    recognized_name = name
                    recognized = True
                    break

            if not recognized:
                cv2.rectangle(frame, (x, y), (x + w, y + h), (0, 255, 0), 2)
                face_file = os.path.join(face_dir, f"face_{len(os.listdir(face_dir)) + 1}.jpg")
                cv2.imwrite(face_file, face_image)
                play_audio("Please tell me the name of this person.")
                person_name = get_voice_input()
                if person_name:
                    person_name = person_name.replace(" ", "_")  # Replace spaces with underscores
                    os.rename(face_file, os.path.join(face_dir, f"{person_name}.jpg"))
                    known_faces[person_name] = face_image
                    with open("known_faces.pkl", "wb") as f:
                        pickle.dump(known_faces, f)
                    play_audio(f"Face saved as {person_name}.")
                    recognized_name = person_name
                else:
                    play_audio("Failed to get the person's name. Please try again.")
                    print("Failed to get the person's name. Please try again.")
                    os.remove(face_file)  # Remove the file if no name is provided
            if recognized_name:
                break

        cv2.imshow('Video', frame)

        if cv2.waitKey(1) & 0xFF == ord('q') or recognized_name:
            break

    video_capture.release()
    cv2.destroyAllWindows()

    return recognized_name


def create_document(user_input, response):
    document = Document()
    document.add_heading(user_input, 0)

    document.add_heading('User Input:', level=1)
    document.add_paragraph(user_input)

    document.add_heading('Gemini Response:', level=1)
    document.add_paragraph(response)

    # Save the document with a unique name
    document.save(user_input)
    return f"Document '{user_input}' created successfully."


# Function to dynamically add new features
def add_feature(feature_description):
    try:
        response = chat_session.send_message(
            f"Create a Python function for the following feature: {feature_description}")
        new_feature_code = response.text
        exec(new_feature_code, globals())
        return "New feature added successfully."
    except Exception as e:
        print(f"Error: {e}")
        return "An error occurred while adding the new feature."


# Function to extract data from Google search results
def extract_google_data(search_query):
    url = f"https://www.googleapis.com/customsearch/v1?q={search_query}&key={GOOGLE_API_KEY}&cx={CSE_ID}"
    response = requests.get(url)
    data = response.json()
    search_results = []
    for item in data.get('items', []):
        result_info = {
            'title': item['title'],
            'link': item['link'],
            'snippet': item['snippet'],
            'seo': item['seo']
        }
        search_results.append(result_info)
    return search_results[:5]


def liveimagesofEarth():
    url = f'https://api.nasa.gov/EPIC/api/natural/images?api_key={NASA_API_KEY}'

    # Directory to save images
    save_dir = 'epic_images'
    if not os.path.exists(save_dir):
        os.makedirs(save_dir)

    # Make the request
    response = requests.get(url)

    # Check if the request was successful
    if response.status_code == 200:
        epic_data = response.json()
        print("Downloading EPIC Images of Earth:")
        for image_info in epic_data:
            image_date = image_info['date'].split(" ")[0].replace("-", "/")
            image_name = image_info['image']
            image_url = f'https://epic.gsfc.nasa.gov/archive/natural/{image_date}/png/{image_name}.png'

            # Download the image
            img_response = requests.get(image_url)

            if img_response.status_code == 200:
                # Save the image to the specified directory
                image_path = os.path.join(save_dir, f"{image_name}.png")
                with open(image_path, 'wb') as f:
                    f.write(img_response.content)
                print(f"Saved: {image_path}")
            else:
                print(f"Failed to download image: {image_url}")
    else:
        print(f"Failed to fetch EPIC data: {response.status_code}")


def apod():
    url = f"https://api.nasa.gov/planetary/apod?api_key={NASA_API_KEY}"

    response = requests.get(NASA_API_KEY)
    if response.status_code == 200:
        data = response.json()

        # Check if the media type is an image
        if data['media_type'] == 'image':
            # Get the image URL from the API response
            image_url = data['url']

            # Fetch the image data
            image_response = requests.get(image_url)
            if image_response.status_code == 200:
                # Convert the image data to a PIL image
                image_data = Image.open(BytesIO(image_response.content))

                # Display the image using Matplotlib
                plt.imshow(image_data)
                plt.axis('off')  # Hide axis
                plt.title(data['title'])  # Set title to APOD title
                plt.show()
            else:
                print("Failed to fetch image from NASA API")
        else:
            print("APOD media type is not an image")
    else:
        print("Failed to fetch astronomy data from NASA API")


# Function to get current location using ipinfo.io
def get_current_location():
    url = "https://ipinfo.io/json"
    response = requests.get(url)
    location_data = response.json()
    loc = location_data['loc'].split(',')
    lat = loc[0]
    lng = loc[1]
    address = location_data['city'] + ", " + location_data['region'] + ", " + location_data['country']
    return lat, lng, address


# Function to get directions using OpenRouteService
def get_directions(destination):
    lat, lng, _ = get_current_location()
    client = openrouteservice.Client(key=ORS_API_KEY)
    coords = ((lng, lat), destination)

    try:
        routes = client.directions(coords)
        steps = routes['routes'][0]['segments'][0]['steps']
        directions = []
        for step in steps:
            directions.append(step['instruction'])
        return directions
    except Exception as e:
        print(f"Error: {e}")
        return get_directions(destination)


def stop_music_listener():
    recognizer = sr.Recognizer()
    while pygame.mixer.music.get_busy():
        with sr.Microphone() as source:
            print("Listening for 'stop' command...")
            audio = recognizer.listen(source)
            try:
                command = recognizer.recognize_google(audio).lower()
                if "stop" in command:
                    pygame.mixer.music.stop()
                    print("Music stopped.")
                    play_audio("Music stopped.")
                    break
            except sr.UnknownValueError:
                pass
            except sr.RequestError:
                print("Could not request results; check your network connection.")
        time.sleep(1)


# Function to get current news
def get_news_by_location(location):
    if location.lower() == "current location":
        _, _, location = get_current_location()

    today_date = datetime.now().strftime("%Y-%m-%d")
    url = f"https://newsapi.org/v2/everything?q={location}&from={today_date}&to={today_date}&sortBy=publishedAt&apiKey={NEWS_API_KEY}"

    response = requests.get(url)
    news_data = response.json()
    headlines = []

    if news_data['status'] == 'ok':
        for article in news_data['articles']:
            headlines.append(article['title'])

    return headlines[:5]


def send_email(recipient, subject, body):
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText

    sender_email = "darkcamperyt007@gmail.com"
    sender_password = "samrudha@125412"

    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = recipient
    msg['Subject'] = subject

    msg.attach(MIMEText(body, 'plain'))

    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(sender_email, sender_password)
        text = msg.as_string()
        server.sendmail(sender_email, recipient, text)
        server.quit()
        return "Email sent successfully."
    except Exception as e:
        print(f"Error: {e}")
        return "Failed to send email."

    # Function to tell a joke


def tell_joke():
    joke = pyjokes.get_joke()
    return joke


def get_nearby_restaurants():
    lat, lng, address = get_current_location()
    url = f"https://maps.googleapis.com/maps/api/place/nearbysearch/json?location={lat},{lng}&radius=1500&type=restaurant&key={GOOGLE_API_KEY}"
    response = requests.get(url)
    places_data = response.json()
    if 'results' not in places_data or len(places_data['results']) == 0:
        return "No nearby restaurants found."

    restaurants = []
    for place in places_data['results'][:5]:  # Limiting to top 5 restaurants
        restaurant_info = {
            'name': place['name'],
            'address': place['vicinity']
        }
        restaurants.append(restaurant_info)

    restaurant_list = "\n".join([f"{r['name']} - {r['address']}" for r in restaurants])
    return f"Here are some nearby restaurants:\n{restaurant_list}"


def control_pc(command):
    if "shutdown" in command:
        os.system("shutdown /s /t 1")  # Shutdown PC
        return "Shutting down the PC."
    elif "restart" in command:
        os.system("shutdown /r /t 1")  # Restart PC
        return "Restarting the PC."
    elif "log off" in command:
        os.system("shutdown /l")  # Log off user
        return "Logging off the user."
    else:
        return "I'm sorry, I couldn't understand that command for controlling the PC."
        # Function to get weather updates


def get_weather(location="current location"):
    if location.lower() == "current location":
        lat, lon, _ = get_current_location()
    else:
        location_url = f"http://api.openweathermap.org/geo/1.0/direct?q={location}&limit=1&appid={WEATHER_API_KEY}"
        location_response = requests.get(location_url).json()
        lat = location_response[0]['lat']
        lon = location_response[0]['lon']

    weather_url = f"https://api.openweathermap.org/data/2.5/weather?lat={lat}&lon={lon}&appid={WEATHER_API_KEY}&units=metric"
    weather_response = requests.get(weather_url).json()

    weather = weather_response['weather'][0]['description']
    temperature = weather_response['main']['temp']
    feels_like = weather_response['main']['feels_like']
    humidity = weather_response['main']['humidity']
    wind_speed = weather_response['wind']['speed']
    weather_info = (
        f"The current weather in {location} is {weather} with a temperature of {temperature}°C, "
        f"feels like {feels_like}°C. The humidity is {humidity}% and the wind speed is {wind_speed} m/s."
    )
    return f"{weather_info}"


# Example functions to fetch nearby data (replace with actual implementations)
def visualize_location(latitude, longitude, address):
    map_center = [latitude, longitude]
    m = folium.Map(location=map_center, zoom_start=15)
    folium.Marker(
        location=map_center,
        popup=f"Current Location: {address}",
        icon=folium.Icon(color='blue', icon='info-sign')
    ).add_to(m)
    m.save('map.html')
    webbrowser.open('map.html')


# Function to handle commands
def handle_command(command):
    command = command.lower()

    if "upload document" in command:
        root = tk.Tk()
        root.withdraw()  # Hide the main window

        file_path = filedialog.askopenfilename()  # Open file picker dialog

        if file_path:
            try:
                with open(file_path, "rb") as file:
                    file_content = file.read()
                response = chat_session.send_file(file_content, filename=file_path)
                return response.text
            except Exception as e:
                print(f"Error uploading document: {e}")
                return "Error uploading document. Please try again."
        else:
            return "No document selected. Please select a document to upload."

        # Other command handling code...

    # Rest of the code remains unchanged...

    if "control PC" in command:
        return "Sure, what would you like to do with your PC?"
    elif "shutdown" in command:
        os.system("shutdown /s /t 1")  # Shutdown PC
        return "Shutting down the PC."
    elif "restart" in command:
        os.system("shutdown /r /t 1")  # Restart PC
        return "Restarting the PC."
    elif "log off" in command:
        os.system("shutdown /l")  # Log off user
        return "Logging off the user."
    elif "images of Earth" in command:
        return liveimagesofEarth()
    elif "exit" in command:
        exit()
    elif "google" in command:
        search_query = command.replace("google", "").strip()
        if search_query:
            results = extract_google_data(search_query)
            return "Here are the top search results: " + ", ".join([result['title'] for result in results])
        else:
            return "You didn't provide a search query."
    elif "open youtube and search" in command:
        search_query = command.replace("open youtube and search", "").strip()
        if search_query:
            url = f"https://www.youtube.com/results?search_query={search_query}"
            webbrowser.open(url)
            return f"Opening YouTube and searching for {search_query}"
        else:
            return "You didn't provide a search query."
    elif "open website" in command:
        website = command.replace("open website", "").strip()
        if website:
            url = f"https://{website}"
            webbrowser.open(url)
            return f"Opening website {website}"
        else:
            return "You didn't provide a website to open."
    elif "add feature" in command:
        feature_description = command.replace("add feature", "").strip()
        if feature_description:
            return add_feature(feature_description)
        else:
            return "You didn't provide a feature description."
    elif "face" in command:
        recognize_faces()
    elif "play" in command:
        partial_name = command.replace("play", "").strip()
        play_music(partial_name)

        # Keep the program running to continue playing music
        while pygame.mixer.music.get_busy():
            pygame.time.Clock().tick(10)
    elif "temperature" in command:
        location = command.replace("temperature in", "").strip()
        return get_weather(location)
    elif "news" in command:
        location = command.replace("news", "").strip()
        return get_news_by_location(location)
    elif "directions to" in command:
        destination = command.replace("directions to", "").strip()
        return get_directions(destination)
    elif "restaurants" in command:
        return get_nearby_restaurants()
    elif "date and time" in command:
        dateAndtime = time.strftime("%d-%m-%Y", time.gmtime())
        play_audio(dateAndtime)

    elif "set alarm of " in command:
        alarm = command.replace("set alarm of", "").strip



    elif "say jokes" in command:
        return tell_joke()
    elif "create document" in command.lower():
        play_audio("Please provide the information to write to the document.")
        info_to_write = get_voice_input()
        if info_to_write:
            response = chat_session.send_message(info_to_write)
            doc_status = create_document(info_to_write, response.text)
            play_audio(doc_status)
        else:
            play_audio("Sorry, I could not understand the information.")

    else:
        response = chat_session.send_message(command)
        print(response.text)
        play_audio(response.text)


# Ensure the environment variable AI_API_KEY is set with your API key
api_key = "AIzaSyBIOrhak4mBQN-FOPfM2Emhu9YOhvaEYq8"
if not api_key:
    raise ValueError("API key not found. Please set the AI_API_KEY environment variable.")

genai.configure(api_key=api_key)

# Create the model and start chat session
generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 64,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}
safety_settings = [
    {
        "category": "HARM_CATEGORY_HARASSMENT",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE",
    },
    {
        "category": "HARM_CATEGORY_HATE_SPEECH",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE",
    },
    {
        "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE",
    },
    {
        "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE",
    },
]

model = genai.GenerativeModel(
    model_name="gemini-1.5-pro-latest",
    safety_settings=safety_settings,
    generation_config=generation_config,
)

chat_session = model.start_chat(
    history=[]
)


def main():
    recognized_name = recognize_faces()
    if recognized_name:
        play_audio(f"Hello, {recognized_name}. I am a simple assistant made by samrudha. How can i assist you today?")
    else:
        play_audio("Face recognition failed. How can I assist you today?")

    while True:
        # Get voice input from the user
        user_message = get_voice_input()
        if user_message:
            # Handle the user command
            response_text = handle_command(user_message)

            # Print response and chat history
            print(response_text)
            print(chat_session.history)

            # Play the response text as audio
            play_audio(response_text)
        else:
            print("Failed to get user input.")

        get_voice_input()


if __name__ == "__main__":
    main()
